"""Loads text out of documents for RAG ingestion.

Supports plain text/code files directly, plus PDF and DOCX extraction.
"""
from __future__ import annotations

from pathlib import Path

TEXT_EXTENSIONS = {".txt", ".md", ".py", ".js", ".html", ".css", ".csv", ".json"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}

SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS


def _load_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def load_text(path: Path) -> str:
    """Dispatches to the right extractor based on file extension."""
    suffix = path.suffix.lower()
    if suffix in PDF_EXTENSIONS:
        return _load_pdf(path)
    if suffix in DOCX_EXTENSIONS:
        return _load_docx(path)
    return _load_plain_text(path)
